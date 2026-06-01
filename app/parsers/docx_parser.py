"""
app/parsers/docx_parser.py — Extracts text and metadata from DOCX files.

Library used: python-docx
A DOCX file is actually a ZIP archive containing XML files. python-docx
unpacks that and gives us a clean Python API to read paragraphs, tables,
headers, and document properties.
"""

from pathlib import Path

from docx import Document


def extract_from_docx(file_path: str) -> dict:
    """
    Open a DOCX file and return its text content and metadata.

    Returns a dict with:
      - "text"       : all text joined into a single string
      - "metadata"   : author, title, dates, paragraph count, etc.
      - "paragraphs" : raw list of non-empty paragraph strings
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    if path.suffix.lower() not in (".docx",):
        raise ValueError(f"Expected a .docx file, got: {path.suffix}")

    doc = Document(file_path)

    # ── Metadata ─────────────────────────────────────────────────────────────
    # doc.core_properties maps to the "Document Properties" panel in Word.
    props = doc.core_properties
    metadata = {
        "title": props.title or "Unknown",
        "author": props.author or "Unknown",
        "subject": props.subject or "",
        "description": props.description or "",
        "created": str(props.created) if props.created else "",
        "modified": str(props.modified) if props.modified else "",
        "last_modified_by": props.last_modified_by or "",
        # DOCX doesn't expose page count directly (it's calculated by Word at
        # render time). We expose paragraph count as a useful proxy.
        "paragraph_count": len(doc.paragraphs),
        "file_name": path.name,
        "file_size_kb": round(path.stat().st_size / 1024, 2),
        "file_type": "DOCX",
    }

    # ── Text Extraction ───────────────────────────────────────────────────────
    paragraphs = []

    # doc.paragraphs covers the main body text (headings, normal text, lists…)
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    # DOCX tables store text in cells — we must read them separately
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    full_text = "\n\n".join(paragraphs)

    return {
        "text": full_text,
        "metadata": metadata,
        "paragraphs": paragraphs,
    }
